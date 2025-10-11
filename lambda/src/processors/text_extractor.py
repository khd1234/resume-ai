"""
Text extraction processor for resume files.

Story 2: Text Extraction from Resume Files - Implementation
Extracts clean text content from PDF and DOCX files for AI analysis.
"""

import logging
import io
import re
from typing import Dict, Any, Optional, List
from datetime import datetime

# PDF extraction libraries
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX extraction library
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from exceptions import TextExtractionError

logger = logging.getLogger(__name__)

class TextExtractor:
    """
    Handles text extraction from PDF and DOCX files.
    
    Implements Story 2 requirements:
    - Clean text extraction from PDF and DOCX files
    - Structure preservation for section identification
    - Error handling for corrupted/password-protected files
    - Text cleaning and formatting
    - Maximum file size validation
    """
    
    def __init__(self):
        """Initialize the text extractor with library availability checks."""
        self.pdf_available = PDF_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        if not PDF_AVAILABLE:
            logger.warning("PDF extraction libraries not available (PyPDF2, pdfplumber)")
        if not DOCX_AVAILABLE:
            logger.warning("DOCX extraction library not available (python-docx)")
    
    def extract_text(self, file_content: bytes, file_type: str, s3_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract text content from resume file.
        
        Args:
            file_content: Raw file content as bytes
            file_type: Type of file (pdf, docx)
            s3_info: S3 file information
            
        Returns:
            Dict with extracted text and metadata
            
        Raises:
            TextExtractionError: If text extraction fails
        """
        start_time = datetime.utcnow()
        file_key = s3_info.get('object_key', 'unknown')
        
        try:
            logger.info(f"Starting text extraction for {file_key} (type: {file_type})")
            
            # Validate file size
            file_size = len(file_content)
            max_size = 2 * 1024 * 1024  # 2MB limit
            if file_size > max_size:
                raise TextExtractionError(
                    f"File size {file_size} bytes exceeds maximum {max_size} bytes for text extraction",
                    file_key=file_key,
                    file_type=file_type
                )
            
            if file_size == 0:
                raise TextExtractionError(
                    "Cannot extract text from empty file",
                    file_key=file_key,
                    file_type=file_type
                )
            
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                result = self._extract_from_pdf(file_content, s3_info)
            elif file_type.lower() == 'docx':
                result = self._extract_from_docx(file_content, s3_info)
            else:
                raise TextExtractionError(
                    f"Unsupported file type for extraction: {file_type}",
                    file_key=file_key,
                    file_type=file_type
                )
            
            # Add extraction timing
            extraction_time = (datetime.utcnow() - start_time).total_seconds()
            result['extraction_duration_seconds'] = round(extraction_time, 3)
            result['file_size_bytes'] = file_size
            
            logger.info(f"Successfully extracted {len(result.get('extracted_text', ''))} characters from {file_key} in {extraction_time:.3f}s")
            return result
            
        except TextExtractionError:
            # Re-raise TextExtractionError as-is
            raise
        except Exception as e:
            logger.error(f"Text extraction failed for {file_key}: {str(e)}", exc_info=True)
            raise TextExtractionError(
                f"Failed to extract text from {file_type}: {str(e)}",
                file_key=file_key,
                file_type=file_type
            )
    
    def _extract_from_pdf(self, file_content: bytes, s3_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract text from PDF file using multiple extraction methods.
        
        Story 2 Implementation:
        - Uses pdfplumber as primary method (better formatting preservation)
        - Falls back to PyPDF2 for compatibility
        - Handles password-protected PDFs gracefully
        - Preserves text structure and layout
        - Detects and reports potential issues
        
        Args:
            file_content: PDF file content as bytes
            s3_info: S3 file information
            
        Returns:
            Dict with extracted text and metadata
        """
        file_key = s3_info.get('object_key', 'unknown')
        warnings = []
        
        if not self.pdf_available:
            raise TextExtractionError(
                "PDF extraction libraries not available",
                file_key=file_key,
                file_type='pdf'
            )
        
        try:
            # Method 1: Try pdfplumber first (better text layout preservation)
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text_parts = []
                    pages_processed = 0
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text.strip())
                                pages_processed += 1
                            else:
                                warnings.append(f"No text found on page {page_num}")
                        except Exception as e:
                            warnings.append(f"Failed to extract text from page {page_num}: {str(e)}")
                            logger.warning(f"PDF page {page_num} extraction failed for {file_key}: {str(e)}")
                    
                    extracted_text = '\n\n'.join(text_parts)
                    
                    if extracted_text.strip():
                        return self._create_pdf_result(
                            extracted_text, pages_processed, len(pdf.pages), 
                            'pdfplumber', warnings, file_key
                        )
                    else:
                        warnings.append("pdfplumber extracted no text")
                        
            except Exception as e:
                warnings.append(f"pdfplumber extraction failed: {str(e)}")
                logger.warning(f"pdfplumber failed for {file_key}: {str(e)}")
            
            # Method 2: Fallback to PyPDF2
            try:
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise TextExtractionError(
                        "Cannot extract text from password-protected PDF",
                        file_key=file_key,
                        file_type='pdf'
                    )
                
                text_parts = []
                pages_processed = 0
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text.strip())
                            pages_processed += 1
                        else:
                            warnings.append(f"No text found on page {page_num} (PyPDF2)")
                    except Exception as e:
                        warnings.append(f"PyPDF2 failed on page {page_num}: {str(e)}")
                        logger.warning(f"PyPDF2 page {page_num} extraction failed for {file_key}: {str(e)}")
                
                extracted_text = '\n\n'.join(text_parts)
                
                if extracted_text.strip():
                    return self._create_pdf_result(
                        extracted_text, pages_processed, len(pdf_reader.pages),
                        'PyPDF2_fallback', warnings, file_key
                    )
                
            except TextExtractionError:
                raise
            except Exception as e:
                warnings.append(f"PyPDF2 extraction failed: {str(e)}")
                logger.error(f"PyPDF2 failed for {file_key}: {str(e)}")
            
            # If both methods failed
            raise TextExtractionError(
                f"All PDF extraction methods failed. Warnings: {'; '.join(warnings)}",
                file_key=file_key,
                file_type='pdf'
            )
            
        except TextExtractionError:
            raise
        except Exception as e:
            raise TextExtractionError(
                f"PDF extraction error: {str(e)}",
                file_key=file_key,
                file_type='pdf'
            )
    
    def _create_pdf_result(self, text: str, pages_processed: int, total_pages: int, 
                          method: str, warnings: List[str], file_key: str) -> Dict[str, Any]:
        """Create standardized PDF extraction result."""
        cleaned_text = self.clean_extracted_text(text)
        sections = self._detect_resume_sections(cleaned_text)
        
        return {
            'extracted_text': cleaned_text,
            'raw_text': text,
            'file_type': 'pdf',
            'extraction_method': method,
            'text_length': len(cleaned_text),
            'pages_processed': pages_processed,
            'total_pages': total_pages,
            'sections_detected': sections,
            'extraction_time': datetime.utcnow().isoformat(),
            'warnings': warnings,
            'success': True
        }
    
    def _extract_from_docx(self, file_content: bytes, s3_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract text from DOCX file with structure preservation.
        
        Story 2 Implementation:
        - Uses python-docx library for reliable extraction
        - Preserves document structure and hierarchy
        - Extracts text from paragraphs, headers, and tables
        - Handles formatting and maintains readability
        - Provides detailed extraction metadata
        
        Args:
            file_content: DOCX file content as bytes
            s3_info: S3 file information
            
        Returns:
            Dict with extracted text and metadata
        """
        file_key = s3_info.get('object_key', 'unknown')
        warnings = []
        
        if not self.docx_available:
            raise TextExtractionError(
                "DOCX extraction library not available (python-docx)",
                file_key=file_key,
                file_type='docx'
            )
        
        try:
            document = Document(io.BytesIO(file_content))
            
            text_parts = []
            paragraphs_processed = 0
            tables_processed = 0
            
            # Extract text from paragraphs
            for para in document.paragraphs:
                try:
                    para_text = para.text.strip()
                    if para_text:
                        text_parts.append(para_text)
                        paragraphs_processed += 1
                except Exception as e:
                    warnings.append(f"Failed to extract paragraph text: {str(e)}")
                    logger.warning(f"DOCX paragraph extraction failed for {file_key}: {str(e)}")
            
            # Extract text from tables
            for table_idx, table in enumerate(document.tables):
                try:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append(f"[Table {table_idx + 1}]\n{table_text}")
                        tables_processed += 1
                except Exception as e:
                    warnings.append(f"Failed to extract table {table_idx + 1}: {str(e)}")
                    logger.warning(f"DOCX table {table_idx + 1} extraction failed for {file_key}: {str(e)}")
            
            # Extract headers and footers
            try:
                header_footer_text = self._extract_headers_footers(document)
                if header_footer_text:
                    text_parts.insert(0, header_footer_text)
            except Exception as e:
                warnings.append(f"Failed to extract headers/footers: {str(e)}")
            
            # Combine all text with proper spacing
            raw_text = '\n\n'.join(text_parts)
            cleaned_text = self.clean_extracted_text(raw_text)
            
            if not cleaned_text.strip():
                raise TextExtractionError(
                    "No extractable text found in DOCX file",
                    file_key=file_key,
                    file_type='docx'
                )
            
            sections = self._detect_resume_sections(cleaned_text)
            
            return {
                'extracted_text': cleaned_text,
                'raw_text': raw_text,
                'file_type': 'docx',
                'extraction_method': 'python-docx',
                'text_length': len(cleaned_text),
                'paragraphs_processed': paragraphs_processed,
                'tables_processed': tables_processed,
                'sections_detected': sections,
                'extraction_time': datetime.utcnow().isoformat(),
                'warnings': warnings,
                'success': True
            }
            
        except TextExtractionError:
            raise
        except Exception as e:
            raise TextExtractionError(
                f"DOCX extraction error: {str(e)}",
                file_key=file_key,
                file_type='docx'
            )
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table with proper formatting."""
        table_rows = []
        
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            
            if row_texts:
                table_rows.append(' | '.join(row_texts))
        
        return '\n'.join(table_rows)
    
    def _extract_headers_footers(self, document) -> str:
        """Extract text from headers and footers."""
        header_footer_parts = []
        
        # Extract from sections (each section can have different headers/footers)
        for section in document.sections:
            try:
                # Header
                if section.header:
                    for para in section.header.paragraphs:
                        header_text = para.text.strip()
                        if header_text and header_text not in header_footer_parts:
                            header_footer_parts.append(f"[Header] {header_text}")
                
                # Footer
                if section.footer:
                    for para in section.footer.paragraphs:
                        footer_text = para.text.strip()
                        if footer_text and footer_text not in header_footer_parts:
                            header_footer_parts.append(f"[Footer] {footer_text}")
            except Exception as e:
                # Headers/footers extraction is best-effort
                logger.debug(f"Header/footer extraction warning: {str(e)}")
        
        return '\n'.join(header_footer_parts)
    
    def clean_extracted_text(self, raw_text: str) -> str:
        """
        Clean and format extracted text for analysis.
        
        Story 2 Implementation:
        - Normalizes whitespace and line breaks
        - Removes formatting artifacts and extra spaces
        - Preserves section structure and readability
        - Handles encoding issues and special characters
        - Maintains important formatting cues
        
        Args:
            raw_text: Raw extracted text
            
        Returns:
            Cleaned and formatted text ready for AI analysis
        """
        if not raw_text:
            return ""
        
        # Step 1: Handle encoding and normalize characters
        try:
            # Ensure proper encoding
            if isinstance(raw_text, bytes):
                raw_text = raw_text.decode('utf-8', errors='ignore')
            
            # Normalize Unicode characters
            import unicodedata
            raw_text = unicodedata.normalize('NFKD', raw_text)
            
        except Exception as e:
            logger.warning(f"Text encoding normalization warning: {str(e)}")
        
        # Step 2: Clean up common formatting artifacts
        text = raw_text
        
        # Remove excessive line breaks (more than 2)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Normalize different types of spaces and tabs
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Clean up bullet points and special characters
        text = re.sub(r'[•·▪▫◦‣⁃]', '•', text)  # Normalize bullet points
        text = re.sub(r'[""''‚„]', '"', text)   # Normalize quotes
        text = re.sub(r'[–—−]', '-', text)      # Normalize dashes
        
        # Step 3: Preserve section breaks but clean up spacing
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Keep non-empty lines
                cleaned_lines.append(line)
            elif cleaned_lines and cleaned_lines[-1]:  # Only add empty line if previous line wasn't empty
                cleaned_lines.append('')
        
        # Step 4: Rejoin with proper spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Step 5: Remove trailing/leading whitespace
        cleaned_text = cleaned_text.strip()
        
        # Step 6: Ensure reasonable text length
        if len(cleaned_text) > 50000:  # 50KB text limit
            logger.warning(f"Text length {len(cleaned_text)} exceeds recommended limit, truncating")
            cleaned_text = cleaned_text[:50000] + "... [Text truncated for processing]"
        
        logger.debug(f"Text cleaning: {len(raw_text)} -> {len(cleaned_text)} characters")
        return cleaned_text
    
    def _detect_resume_sections(self, text: str) -> List[Dict[str, Any]]:
        """
        Detect common resume sections in the extracted text.
        
        Story 2 Enhancement:
        - Identifies typical resume sections (contact, summary, experience, etc.)
        - Provides section boundaries for better AI analysis
        - Uses pattern matching and keyword detection
        - Helps maintain document structure understanding
        
        Args:
            text: Cleaned extracted text
            
        Returns:
            List of detected sections with positions and types
        """
        sections = []
        
        # Common resume section patterns
        section_patterns = {
            'contact': [
                r'contact\s+information',
                r'personal\s+information',
                r'contact\s+details',
                r'phone.*email',
                r'email.*phone'
            ],
            'summary': [
                r'professional\s+summary',
                r'career\s+summary',
                r'executive\s+summary',
                r'profile',
                r'objective',
                r'summary\s+of\s+qualifications'
            ],
            'experience': [
                r'work\s+experience',
                r'professional\s+experience',
                r'employment\s+history',
                r'career\s+history',
                r'experience'
            ],
            'education': [
                r'education',
                r'academic\s+background',
                r'educational\s+qualifications'
            ],
            'skills': [
                r'technical\s+skills',
                r'core\s+competencies',
                r'skills\s+and\s+abilities',
                r'key\s+skills',
                r'skills'
            ],
            'certifications': [
                r'certifications',
                r'certificates',
                r'professional\s+certifications',
                r'licenses'
            ]
        }
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        for section_type, patterns in section_patterns.items():
            for pattern in patterns:
                matches = list(re.finditer(pattern, text_lower, re.IGNORECASE))
                for match in matches:
                    # Find the line number for context
                    char_pos = match.start()
                    line_num = text[:char_pos].count('\n') + 1
                    
                    sections.append({
                        'type': section_type,
                        'pattern': pattern,
                        'position': char_pos,
                        'line_number': line_num,
                        'matched_text': match.group()
                    })
        
        # Sort by position in document
        sections.sort(key=lambda x: x['position'])
        
        logger.debug(f"Detected {len(sections)} resume sections")
        return sections
    
    def get_extraction_capabilities(self) -> Dict[str, bool]:
        """
        Get current extraction capabilities based on available libraries.
        
        Returns:
            Dict indicating which file types can be processed
        """
        return {
            'pdf_extraction': self.pdf_available,
            'docx_extraction': self.docx_available,
            'libraries': {
                'PyPDF2': PDF_AVAILABLE and 'PyPDF2' in globals(),
                'pdfplumber': PDF_AVAILABLE and 'pdfplumber' in globals(),
                'python-docx': DOCX_AVAILABLE
            }
        }